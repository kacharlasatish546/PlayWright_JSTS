from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Reduce margins
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h

def add_horizontal_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '1A3C6E'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

# ── NAME ──
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name_para.add_run('JOHN DOE')
name_run.bold = True
name_run.font.size = Pt(26)
name_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── TITLE ──
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('Senior QA Automation Engineer')
title_run.bold = True
title_run.font.size = Pt(14)
title_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ── CONTACT ──
contact_para = doc.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_para.paragraph_format.space_after = Pt(4)
contact_run = contact_para.add_run('+1 (555) 123-4567  |  john.doe@email.com  |  linkedin.com/in/johndoe  |  Hyderabad, India')
contact_run.font.size = Pt(10)

add_horizontal_line()

# ── PROFESSIONAL SUMMARY ──
add_heading_styled('PROFESSIONAL SUMMARY', level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.add_run(
    'Results-driven Senior QA Automation Engineer with 9+ years of experience in designing, developing, '
    'and maintaining robust test automation frameworks. Expertise in Playwright, Selenium, Cypress, and API '
    'testing using JavaScript/TypeScript and Java. Proven track record in implementing CI/CD pipelines, leading '
    'QA teams, and establishing quality best practices across Agile/Scrum environments. Strong background in web, '
    'mobile, and API test automation with a focus on scalable, maintainable, and reliable test solutions.'
)

add_horizontal_line()

# ── TECHNICAL SKILLS ──
add_heading_styled('TECHNICAL SKILLS', level=2)

skills = [
    ('Automation Tools', 'Playwright, Selenium WebDriver, Cypress, Appium, WebDriverIO, Puppeteer'),
    ('Programming Languages', 'JavaScript, TypeScript, Java, Python'),
    ('Test Frameworks', 'TestNG, JUnit, Mocha, Jest, Cucumber (BDD), pytest'),
    ('API Testing', 'Postman, REST Assured, Axios, Supertest, Swagger'),
    ('CI/CD', 'Jenkins, GitHub Actions, Azure DevOps, GitLab CI, CircleCI'),
    ('Version Control', 'Git, GitHub, Bitbucket, GitLab'),
    ('Cloud & Containers', 'Docker, AWS (EC2, S3, Lambda), BrowserStack, Sauce Labs'),
    ('Databases', 'MySQL, PostgreSQL, MongoDB, Oracle'),
    ('Reporting', 'Allure, ExtentReports, HTML Reports, Playwright HTML Reporter'),
    ('Project Management', 'JIRA, Azure Boards, Confluence, TestRail, Zephyr'),
    ('Other', 'Page Object Model (POM), Data-Driven Testing, Parallel Execution, Cross-Browser Testing, Performance Testing (JMeter, k6)'),
]

table = doc.add_table(rows=len(skills), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light List Accent 1'

for i, (cat, sk) in enumerate(skills):
    row = table.rows[i]
    c1 = row.cells[0]
    c2 = row.cells[1]
    c1_run = c1.paragraphs[0].add_run(cat)
    c1_run.bold = True
    c1_run.font.size = Pt(10)
    c2_run = c2.paragraphs[0].add_run(sk)
    c2_run.font.size = Pt(10)

add_horizontal_line()

# ── PROFESSIONAL EXPERIENCE ──
add_heading_styled('PROFESSIONAL EXPERIENCE', level=2)

# Role 1
p = doc.add_paragraph()
run = p.add_run('Senior QA Automation Engineer')
run.bold = True
run.font.size = Pt(12)
p2 = doc.add_paragraph()
run2 = p2.add_run('ABC Technology Solutions')
run2.bold = True
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p2.add_run('  |  Hyderabad, India  |  Jan 2022 – Present').font.color.rgb = RGBColor(0x77, 0x77, 0x77)
p2.paragraph_format.space_after = Pt(4)

bullets1 = [
    'Architected and built an end-to-end Playwright automation framework (TypeScript) covering 500+ test cases across 3 web applications, reducing manual testing effort by 70%',
    'Led a team of 5 QA engineers, mentoring juniors on best practices in test design, code reviews, and automation patterns',
    'Implemented CI/CD pipelines using GitHub Actions with parallel test execution, achieving test run times under 15 minutes for full regression',
    'Designed API test automation using REST Assured and Supertest for 200+ API endpoints with contract testing and schema validation',
    'Established BDD framework using Cucumber with Playwright, enabling collaboration between QA, Dev, and Business teams',
    'Integrated visual regression testing using Playwright\'s screenshot comparison, catching 30+ UI defects before production releases',
    'Configured cross-browser and cross-device testing using BrowserStack, ensuring compatibility across Chrome, Firefox, Safari, and Edge',
    'Reduced test flakiness from 15% to under 2% through robust retry mechanisms, smart waits, and test isolation strategies',
]
for b in bullets1:
    add_bullet(b)

# Role 2
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('QA Automation Engineer')
run.bold = True
run.font.size = Pt(12)
p2 = doc.add_paragraph()
run2 = p2.add_run('XYZ Software Pvt. Ltd.')
run2.bold = True
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p2.add_run('  |  Bangalore, India  |  Jun 2019 – Dec 2021').font.color.rgb = RGBColor(0x77, 0x77, 0x77)
p2.paragraph_format.space_after = Pt(4)

bullets2 = [
    'Developed and maintained Selenium WebDriver automation framework using Java and TestNG with Page Object Model design pattern',
    'Automated 300+ regression test cases for an e-commerce platform, reducing regression cycle from 5 days to 8 hours',
    'Created data-driven test suites using Apache POI and CSV parsers for testing multiple user scenarios',
    'Built API automation suite using REST Assured for microservices-based architecture covering CRUD operations and authentication flows',
    'Implemented Allure reporting with detailed step-by-step logs, screenshots on failure, and execution history tracking',
    'Set up Jenkins CI/CD pipelines with Selenium Grid for parallel execution across multiple browser configurations',
    'Conducted performance testing using JMeter for load and stress testing of critical business workflows',
    'Participated in Agile ceremonies including sprint planning, daily standups, and retrospectives',
]
for b in bullets2:
    add_bullet(b)

# Role 3
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('QA Automation Engineer')
run.bold = True
run.font.size = Pt(12)
p2 = doc.add_paragraph()
run2 = p2.add_run('PQR InfoTech')
run2.bold = True
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p2.add_run('  |  Pune, India  |  Mar 2017 – May 2019').font.color.rgb = RGBColor(0x77, 0x77, 0x77)
p2.paragraph_format.space_after = Pt(4)

bullets3 = [
    'Built hybrid automation framework using Selenium WebDriver with Java, integrating TestNG and Cucumber for BDD',
    'Automated functional, regression, and smoke test suites for a banking application with 200+ test scenarios',
    'Developed reusable utility libraries for common operations like file handling, database validation, and email verification',
    'Performed database testing using JDBC to validate data integrity across Oracle and MySQL databases',
    'Collaborated with development teams to implement shift-left testing practices, identifying defects 40% earlier in the SDLC',
    'Created and maintained test documentation including test plans, test cases, and defect reports in JIRA and TestRail',
]
for b in bullets3:
    add_bullet(b)

# Role 4
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('QA Analyst / Junior Automation Engineer')
run.bold = True
run.font.size = Pt(12)
p2 = doc.add_paragraph()
run2 = p2.add_run('LMN Solutions')
run2.bold = True
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p2.add_run('  |  Hyderabad, India  |  Apr 2015 – Feb 2017').font.color.rgb = RGBColor(0x77, 0x77, 0x77)
p2.paragraph_format.space_after = Pt(4)

bullets4 = [
    'Executed manual testing for web and mobile applications across multiple domains including healthcare and retail',
    'Transitioned from manual to automation testing, developing initial Selenium WebDriver scripts using Java',
    'Created and maintained test cases and test scripts for functional, integration, and system testing',
    'Performed cross-browser testing and mobile responsive testing using BrowserStack',
    'Logged and tracked defects using JIRA with detailed reproduction steps and severity classification',
    'Participated in requirement analysis and test estimation activities',
]
for b in bullets4:
    add_bullet(b)

add_horizontal_line()

# ── KEY PROJECTS ──
add_heading_styled('KEY PROJECTS', level=2)

projects = [
    ('E-Commerce Platform Test Automation (Playwright + TypeScript)', [
        'Built comprehensive test suite covering user registration, product search, cart management, checkout, and payment flows',
        'Implemented visual testing, accessibility testing (axe-core), and performance metrics collection',
        'Achieved 95% automation coverage with zero false positives in CI pipeline',
    ]),
    ('Banking Application API Automation (REST Assured + Java)', [
        'Automated 150+ API test cases covering account management, transactions, KYC, and authentication modules',
        'Implemented contract testing and response schema validation using JSON Schema Validator',
        'Integrated with Jenkins pipeline for nightly regression runs with Allure reporting',
    ]),
    ('Mobile App Automation (Appium + JavaScript)', [
        'Developed cross-platform test automation for iOS and Android banking application',
        'Automated 100+ test cases covering login, fund transfer, bill payments, and notifications',
        'Configured cloud execution on BrowserStack App Automate for device farm testing',
    ]),
]

for title, items in projects:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    for item in items:
        add_bullet(item)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_horizontal_line()

# ── EDUCATION ──
add_heading_styled('EDUCATION', level=2)
p = doc.add_paragraph()
run = p.add_run('Bachelor of Technology (B.Tech) in Computer Science')
run.bold = True
run.font.size = Pt(11)
p2 = doc.add_paragraph()
run2 = p2.add_run('JNTU University')
run2.bold = True
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
p2.add_run('  |  Hyderabad, India  |  2011 – 2015').font.color.rgb = RGBColor(0x77, 0x77, 0x77)

add_horizontal_line()

# ── CERTIFICATIONS ──
add_heading_styled('CERTIFICATIONS', level=2)
certs = [
    'ISTQB Certified Tester – Foundation Level (2016)',
    'ISTQB Certified Tester – Advanced Level (Test Automation Engineer) (2019)',
    'AWS Certified Cloud Practitioner (2022)',
    'Certified Selenium Professional (2018)',
]
for c in certs:
    add_bullet(c)

add_horizontal_line()

# ── ACHIEVEMENTS ──
add_heading_styled('ACHIEVEMENTS', level=2)
achievements = [
    'Received "Star Performer" award at ABC Technology Solutions for driving 70% reduction in manual testing effort (2023)',
    'Published internal knowledge base articles on Playwright best practices adopted by 50+ engineers across the organization',
    'Speaker at internal tech talks on "Migrating from Selenium to Playwright" and "Flaky Test Detection and Prevention"',
    'Contributed to open-source Playwright utility libraries on GitHub',
]
for a in achievements:
    add_bullet(a)

# Save
doc.save(r'd:\Satish\Playwright\resume.docx')
print('resume.docx created successfully!')
