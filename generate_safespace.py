from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── STYLES & MARGINS ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

DARK_BLUE = RGBColor(0x0B, 0x2D, 0x5E)
MED_BLUE = RGBColor(0x1A, 0x5C, 0x9E)
GRAY = RGBColor(0x55, 0x55, 0x55)

def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_horizontal_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): '0B2D5E'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = DARK_BLUE
        r.font.size = Pt(16)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    return h

def heading2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = MED_BLUE
        r.font.size = Pt(13)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    return h

def heading3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = MED_BLUE
        r.font.size = Pt(11)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(3)
    return h

def para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.3 + level * 0.35)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def bold_bullet(label, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.3 + level * 0.35)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p

def info_table(rows_data):
    t = doc.add_table(rows=len(rows_data), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows_data):
        c1, c2 = t.rows[i].cells
        set_cell_shading(c1, 'E8EEF7')
        r1 = c1.paragraphs[0].add_run(k)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = c2.paragraphs[0].add_run(v)
        r2.font.size = Pt(10)
    return t


# ═══════════════════════════════════════════════════════
# COVER / TITLE PAGE
# ═══════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('SAFE SPACE for School')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Functional Requirements Document (FRD)')
r.font.size = Pt(16)
r.font.color.rgb = GRAY

doc.add_paragraph()
add_horizontal_line()
doc.add_paragraph()

info_table([
    ('Document Type', 'Functional Requirements Document'),
    ('Project Name', 'SAFE SPACE – School Security Platform'),
    ('Version', '1.0'),
    ('Status', 'Draft'),
    ('Prepared By', '[Author Name]'),
    ('Reviewed By', '[Reviewer Name]'),
    ('Approved By', '[Approver Name]'),
    ('Date', 'April 2026'),
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# DOCUMENT CONTROL
# ═══════════════════════════════════════════════════════
heading1('Document Control')

heading2('Revision History')
t = doc.add_table(rows=2, cols=5)
t.style = 'Light List Accent 1'
headers = ['Version', 'Date', 'Author', 'Change Description', 'Status']
for i, h in enumerate(headers):
    r = t.rows[0].cells[i].paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9)
data = ['1.0', 'April 2026', '[Author]', 'Initial document creation', 'Draft']
for i, d in enumerate(data):
    t.rows[1].cells[i].paragraphs[0].add_run(d).font.size = Pt(9)

doc.add_paragraph()

heading2('Distribution List')
t2 = doc.add_table(rows=4, cols=3)
t2.style = 'Light List Accent 1'
for i, h in enumerate(['Role', 'Name', 'Access Level']):
    t2.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
dist_data = [
    ('Project Manager', '[Name]', 'Full Access'),
    ('QA Lead', '[Name]', 'Full Access'),
    ('Development Lead', '[Name]', 'Read Only'),
]
for ri, (role, name, access) in enumerate(dist_data, 1):
    for ci, val in enumerate([role, name, access]):
        t2.rows[ri].cells[ci].paragraphs[0].add_run(val).font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════
heading1('Table of Contents')
toc_items = [
    '1. Login Page',
    '2. Dashboard',
    '   2.1 Type of Visitors',
    '   2.2 Threat Matrix',
    '   2.3 Live Feed',
    '   2.4 Alerts Triggered',
    '   2.5 Alert Console',
    '      2.5.1 Active Console',
    '      2.5.2 Archived Tab',
    '      2.5.3 Settings',
    '3. Analytics and Investigation',
    '   3.1 Analytics',
    '   3.2 Investigations (Roll Call, POI, OOI, Camera, Agent Sam)',
    '4. Live Feed',
    '5. Rules and Alerts',
    '   5.1 Rules Engine',
    '   5.2 Alerts',
    '   5.3 Human in Loop Console',
    '6. Configurations',
    '   6.1 School/Space Configuration',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item)
    r.font.size = Pt(10.5)
    r.font.color.rgb = MED_BLUE

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# MODULE 1 – LOGIN PAGE
# ═══════════════════════════════════════════════════════
heading1('1. Login Page')

para('Below is the URL for logging into the QA Safe Space system:')
p = doc.add_paragraph()
r = p.add_run('SafeSpace | Schools')
r.font.color.rgb = MED_BLUE
r.underline = True
r.font.size = Pt(10.5)

para('Below are the requirements for logging into the QA environment:')

bold_bullet('Username: ', 'The user should be entering an already registered email ID in place of user name.')
bold_bullet('Password: ', 'For the first user, the user would be getting an OTP to their mobile number and once they enter the mobile number, the person has to create a new password and confirm the password.')
bold_bullet('Forgot Password: ', 'The user can click on forgot password in case user had forgotten the password and the system will redirect for forgot password screen where the user can enter the code sent to their email id. Once that is entered, user can reset the password.')
bold_bullet('OTP Login: ', 'For every login, the user will be receiving a unique OTP which they have to enter to login to the system.')

para('After entering the above, the user can login into the SAFE SPACE school portal.')

add_horizontal_line()

# ═══════════════════════════════════════════════════════
# MODULE 2 – DASHBOARD
# ═══════════════════════════════════════════════════════
heading1('2. Dashboard')
para('The Dashboard displays various information as listed below:')

bullet('Types of visitors with number')
bullet('Threat Matrix')
bullet('Live Feed')
bullet('Alerts Triggered (Time range)')
bullet('Quick Actions')

heading2('2.1 Type of Visitors')
bold_bullet('Invited Visitors: ', 'Here the tab is clickable and upon the click, it displays the list details of each visitor such as a) Name, b) ID, c) Email ID, d) Mobile Number, e) Status.')
bold_bullet('Registered Users: ', 'Here the tab is clickable and upon the click, it displays the list details of each registered user such as a) Name, b) ID, c) Email ID, d) Mobile Number, e) Status.')
bold_bullet('Inactive Users: ', 'Here the tab is clickable and upon the click, it displays the list details of each Inactive user such as a) Name, b) ID, c) Email ID, d) Mobile Number, e) Status.')

heading2('2.2 Threat Matrix')
para('This displays the threat detected and upon hover, it displays details such as:')
bullet('Location', 1)
bullet('Camera Name', 1)
bullet('School Name', 1)
bullet('Threat Type', 1)
bullet('Time Stamp', 1)
bullet('Recommended Action', 1)
bullet('Image (opens up the image of the threat)', 1)

heading2('2.3 Live Feed')
para('This displays all the cameras that are pinned.')

heading2('2.4 Alerts Triggered')
para('Here the user can select the date range within which he needs the data.')

heading2('2.5 Alert Console')
para('Upon clicking on Alert Console, it takes the user to a different screen which displays the user\'s clock. This displays the below tabs:')
bullet('Active Console')
bullet('Archived Tab')
bullet('Settings')

heading3('2.5.1 Active Console')
para('This displays the below tabs: severity, tracker, call to action/human in loop, alerts timestamp, status, description. Except tracker and description, rest others are non-clickable. Also displays various filters as below:')

bold_bullet('Severity: ', 'Drop-down menu containing the level of severity — Low, High, Medium, Critical and also All in a single view. Based on selection of this filter, the system displays the required details.')
bold_bullet('Anomaly Type: ', 'Drop-down menu which displays Monitor and track onsite vehicle time log, watch list person detected, weapon detected. Based on selection of this filter, the system displays the required details.')
bold_bullet('Space: ', 'This is also a dropdown which upon selection, displays from which side of the camera the user wants to check the details.')
bold_bullet('Action Taken: ', 'This is a dropdown which displays the details of the action taken for a particular incident that had taken place such as was the information sent to 911, Email, SMS or Rapid SOS.')
bold_bullet('Search: ', 'General search where the user can search based on his knowledge and experience.')

para('Tracker Field:', bold=True)
para('This is a clickable field which upon click displays:')

p = doc.add_paragraph()
r = p.add_run('Alert Details:')
r.bold = True
r.font.size = Pt(10.5)
bullet('Event timestamp of the alert', 1)
bullet('Description, Space details, Severity, Status', 1)
bullet('Change status dropdown (reviewed, false positive, completed notes, audited notes, action, action taken)', 1)
bullet('Camera footage — can be expanded', 1)
bullet('Action taken comments', 1)
bullet('Save change if any or Cancel', 1)

p = doc.add_paragraph()
r = p.add_run('Track Incident Tab — contains two sub-tabs:')
r.bold = True
r.font.size = Pt(10.5)
bold_bullet('Action Taken Tab: ', 'Displays the Alert ID, severity, location, response timeline / Audit trail / log.', 1)
bold_bullet('Notification Sent Tab: ', 'Displays the list of users to whom the notifications have been sent along with date, timestamp, and status of notification (delivered, queued).', 1)

para('Description Field:', bold=True)
para('This is a clickable field which upon click displays:')

p = doc.add_paragraph()
r = p.add_run('Alert Details:')
r.bold = True
r.font.size = Pt(10.5)
bullet('Event timestamp of the alert', 1)
bullet('Description, Space details, Severity, Status', 1)
bullet('Change status dropdown (reviewed, false positive, completed notes, audited notes, action, action taken)', 1)
bullet('Camera footage — can be expanded', 1)
bullet('Action taken comments', 1)
bullet('Save change if any or Cancel', 1)

p = doc.add_paragraph()
r = p.add_run('Track Incident Tab — contains two sub-tabs:')
r.bold = True
r.font.size = Pt(10.5)
bold_bullet('Action Taken Tab: ', 'Displays the Alert ID, severity, location, response timeline / Audit trail / log.', 1)
bold_bullet('Notification Sent Tab: ', 'Displays the list of users to whom the notifications have been sent along with date, timestamp, and status of notification (delivered, queued).', 1)

heading3('2.5.2 Archived Tab')
para('This displays the below tabs: severity, timestamp, status and description (clickable). This also displays below filters:')

bold_bullet('Severity: ', 'Drop-down menu containing the level of severity — Low, High, Medium, Critical and also All in a single view. Based on selection of this filter, the system displays the required details.')
bold_bullet('Anomaly Type: ', 'Drop-down menu which displays Monitor and track onsite vehicle time log, watch list person detected, weapon detected. Based on selection of this filter, the system displays the required details.')
bold_bullet('Space: ', 'This is also a dropdown which upon selection, displays from which side of the camera the user wants to check the details.')
bold_bullet('Action Taken: ', 'This is a dropdown which displays the details of the action taken for a particular incident that had taken place such as was the information sent to 911, Email, SMS or Rapid SOS.')
bold_bullet('Search: ', 'General search where the user can search based on his knowledge and experience.')

para('Description Field:', bold=True)
para('This is a clickable field which upon click displays:')

p = doc.add_paragraph()
r = p.add_run('Alert Details:')
r.bold = True
r.font.size = Pt(10.5)
bullet('Event timestamp of the alert', 1)
bullet('Description, Space details, Severity, Status', 1)
bullet('Change status dropdown (reviewed, false positive, completed notes, audited notes, action, action taken)', 1)
bullet('Camera footage — can be expanded', 1)
bullet('Action taken comments', 1)
bullet('Save change if any or Cancel', 1)

p = doc.add_paragraph()
r = p.add_run('Track Incident Tab — contains two sub-tabs:')
r.bold = True
r.font.size = Pt(10.5)
bold_bullet('Action Taken Tab: ', 'Displays the Alert ID, severity, location, response timeline / Audit trail / log.', 1)
bold_bullet('Notification Sent Tab: ', 'Displays the list of users to whom the notifications have been sent along with date, timestamp, and status of notification (delivered, queued).', 1)

heading3('2.5.3 Settings')
para('These are the settings available based on the roles assigned to a user. Below are the settings available for an Admin:')

para('Admin Settings:', bold=True)
para('The admin can sort the settings with the dropdown menu such as: By severity and By created date.')

para('Escalation Settings:', bold=True)
para('Here the admin can create rules for escalation based on the below factors:')

bold_bullet('Rules: ', 'Rules are the type of escalations that the admin wants to see in an alert (Brandished guns & knives, guns and knives, vehicle time log, watchlist person detected).')
para('When Brandished guns & knives and watchlist person detected are selected, the rule displays: Types of Levels, Escalated to, Pick actions and notifications. Here the Levels that can be added are up to 3 levels. There is also a checkbox (Auto expire in 5 min) — if this checkbox is selected, the escalation rule will get expired in 5 mins. When the Admin clicks on Save, the rule gets saved.', indent=True)
para('Below the rules tab, this displays the list of rules created.', indent=True)

para('Alert Assignment:', bold=True)
para('Here, the admin can create a rule to whom the alerts can be sent:')
bullet('Brandished Guns & Knives — Super Admin and Security Officer', 1)
bullet('Guns & Knives 123 — NA', 1)
bullet('Vehicle Time Log — Security Admin', 1)
bullet('Watchlist Person Detected — Super Admin and Security Officer', 1)

para('Assign Camera Links:', bold=True)
para('When the Admin clicks on the link, the screen redirects to another screen where the admin can configure the camera such as adding a new camera, select school name and space from the filters. This page also lists down all the details of cameras as camera name, manufacturer, camera type, last inspected on, status and actions taken.')

p = doc.add_paragraph()
r = p.add_run('Adding New Camera — Required Fields:')
r.bold = True
r.font.size = Pt(10.5)

cam_fields = [
    ('1.', 'Select School (dropdown)'),
    ('2.', 'Select Space (dropdown)'),
    ('3.', 'Enter Camera Name (text box)'),
    ('4.', 'Select Manufacturer (dropdown)'),
    ('5.', 'Select Camera Type (dropdown)'),
    ('6.', 'Select Model (dropdown)'),
    ('7.', 'Enter Description (text box)'),
    ('8.', 'Select SafeCore Box ID (dropdown)'),
    ('9.', 'Last Inspected Date (date range)'),
    ('10.', 'Username (text box)'),
    ('11.', 'Password (text box)'),
    ('12.', 'Enter Port (text box)'),
    ('13.', 'Enter IP Address (text box)'),
    ('14.', 'Select Status (dropdown)'),
    ('15.', 'Live Feed URL (text box)'),
    ('16.', 'Enter RTSP URL (text box)'),
]
t = doc.add_table(rows=len(cam_fields)+1, cols=3)
t.style = 'Light List Accent 1'
for ci, h in enumerate(['#', 'Field Name', 'Input Type']):
    t.rows[0].cells[ci].paragraphs[0].add_run(h).bold = True
for ri, (num, field) in enumerate(cam_fields, 1):
    t.rows[ri].cells[0].paragraphs[0].add_run(num).font.size = Pt(9)
    fname = field.split('(')[0].strip()
    ftype = '(' + field.split('(')[1] if '(' in field else ''
    t.rows[ri].cells[1].paragraphs[0].add_run(fname).font.size = Pt(9)
    t.rows[ri].cells[2].paragraphs[0].add_run(ftype).font.size = Pt(9)

para('Once the Admin enters all the above details and saves, the details of the camera get saved in the database.', indent=True)

add_horizontal_line()

add_horizontal_line()

# ═══════════════════════════════════════════════════════
# MODULE 3 – ANALYTICS AND INVESTIGATION
# ═══════════════════════════════════════════════════════
heading1('3. Analytics and Investigation')
para('Under this tab, there are 2 sub-tabs. This tab basically displays the analysis of the recorded videos. Below are the two tabs available:')
bullet('Analytics')
bullet('Investigations')

heading2('3.1 Analytics')
para('This displays the threat detected and upon hover, it displays details such as:')
bullet('Location', 1)
bullet('Camera Name', 1)
bullet('School Name', 1)
bullet('Threat Type', 1)
bullet('Time Stamp', 1)
bullet('Recommended Action', 1)
bullet('Image (opens up the image of the threat)', 1)

heading2('3.2 Investigations')
para('Under investigations, the user finds out the reason/person for the issue caused. There are 3 sub-tabs:')
bullet('Roll Call')
bullet('Investigations')
bullet('Agent Sam')

heading3('3.2.1 Roll Call')
bullet('Under roll call, the user can find out the category of person entering the school premises when selects the category type (dropdown) and search for a specific person in the search dropdown and can validate with the photo displayed.')
bullet('When the user selects the date range and wanted to see any particular visitor\'s data with a particular date range, this displays that.')
bullet('Similarly, the user can simply search the details of a particular person from doing a simple search.')

heading3('3.2.2 Investigations — POI (Person of Interest)')
para('Here the user can select the Person of Interest (POI) categories:')

poi_items = [
    'Gender', 'Height', 'Mask', 'Hair Colour', 'Beard', 'Eye Wear',
    'Clothing Wear', 'Upper Clothing Type', 'Lower Clothing Type',
    'Upper Clothing Colour', 'Lower Clothing Colour', 'Accessories', 'Headwear'
]
t = doc.add_table(rows=1, cols=4)
t.style = 'Light List Accent 1'
for ci, h in enumerate(['#', 'Filter Category', 'Input Type', 'Required']):
    t.rows[0].cells[ci].paragraphs[0].add_run(h).bold = True
for idx, item in enumerate(poi_items, 1):
    row = t.add_row()
    row.cells[0].paragraphs[0].add_run(str(idx)).font.size = Pt(9)
    row.cells[1].paragraphs[0].add_run(item).font.size = Pt(9)
    row.cells[2].paragraphs[0].add_run('Dropdown').font.size = Pt(9)
    row.cells[3].paragraphs[0].add_run('Optional').font.size = Pt(9)

doc.add_paragraph()
para('All categories have sub-categories in the dropdown menu. The user can also select the date range and once the user clicks on Investigate, the system displays all the videos containing the filters applied. There is also a feature to start a new chat where the system will refresh the filters and the user can select new filters again. When the user selects one or more filters, the system will display the results accordingly.')

heading3('3.2.3 Investigations — OOI (Object of Interest)')
para('Here the user can search based on the Object of Interest:')
bold_bullet('License Plate Registration: ', 'Text box input')
bold_bullet('License Plate Stop Word: ', 'Text box input')
bold_bullet('Weapon: ', 'Dropdown selection')
para('When the user selects one or more filters and clicks on Investigate, the system will display the results accordingly.')

heading3('3.2.4 Investigations — Camera')
para('Here the user can select:')
bold_bullet('School Name: ', 'Dropdown')
bold_bullet('Space: ', 'Dropdown')
bold_bullet('Camera Name: ', 'Dropdown')
bold_bullet('Date Range: ', 'Date picker')
para('When the user selects one or more filters and clicks on Investigate, the system will display the results accordingly.')

heading3('3.2.5 Agent Sam')
para('This is a BOT where it processes the queries asked and displays the query results accordingly. This BOT could take 24 hours to process any images added to the queries. The user can try the next day if the search results are failing.')

add_horizontal_line()

# ═══════════════════════════════════════════════════════
# MODULE 4 – LIVE FEED
# ═══════════════════════════════════════════════════════
heading1('4. Live Feed')
para('Here, the user can view the live videos from the cameras. The user can also filter from which place and which side of the place they want to view the results. On the top right side of the page, the user can select the number of cameras to be displayed and the pattern in which the cameras should be displayed.')

add_horizontal_line()

# ═══════════════════════════════════════════════════════
# MODULE 5 – RULES AND ALERTS
# ═══════════════════════════════════════════════════════
heading1('5. Rules and Alerts')
para('Here the admin can configure rules for alerts, check if the users are receiving proper alerts, analyse the alerts and also manage human in loop.')

heading2('5.1 Rules Engine')
para('The screen displays the list of already existing rules with the following details:')

t = doc.add_table(rows=1, cols=6)
t.style = 'Light List Accent 1'
for ci, h in enumerate(['Rule Name', 'Rule Description', 'Rule Level', 'Anomaly', 'Severity', 'Action']):
    t.rows[0].cells[ci].paragraphs[0].add_run(h).bold = True
row = t.add_row()
for ci, v in enumerate(['[Name]', '[Description]', '[Level]', '[Type]', '[Level]', 'Edit / Delete']):
    row.cells[ci].paragraphs[0].add_run(v).font.size = Pt(9)

heading3('5.1.1 Creation of New Rule')
para('When the admin wants to create any new rule and clicks on the Add button, a new screen pops up where admin needs to provide the below details:')

p = doc.add_paragraph()
r = p.add_run('Rule Level Options:')
r.bold = True
r.font.size = Pt(10.5)

rule_levels = [
    ('Global: ', 'When the Admin selects Global, the rule is applicable globally. Admin has to enter rule name in the text box provided, provide the description, choose anomaly (dropdown), choose severity (dropdown), can also choose until when the rule is applicable (date range) and save the rule.'),
    ('School: ', 'When the admin selects School, the rule is applicable to that particular school in the location selected from the dropdown. Once the admin selects this, the information related to that particular school gets populated. Similar to global, once the user enters the rule name, admin has to provide the description, choose anomaly (dropdown), choose severity (dropdown), can also choose until when the rule is applicable (date range) and save the rule.'),
    ('Space: ', 'When the admin selects the location from the dropdown and selects the space, the description shows alert configuration details as well. Once the admin selects this, the information related to that particular school gets populated. Similar to global, once the user enters the rule name, admin has to provide the description, choose anomaly (dropdown), choose severity (dropdown), can also choose until when the rule is applicable (date range) and save the rule.'),
    ('Telemetry: ', 'When the admin selects the location from the dropdown and selects the space, the description shows alert configuration details as well. Once the admin selects this, the information related to that particular school gets populated. Similar to global, once the user enters the rule name, admin has to provide the description, choose anomaly (dropdown), choose severity (dropdown), can also choose until when the rule is applicable (date range) and save the rule.'),
]
for label, text in rule_levels:
    bold_bullet(label, text)

heading2('5.2 Alerts')
para('Alert Console: Upon clicking on alert console, it takes the user to a different screen which displays the user\'s clock. This displays the below tabs:')
bullet('Active Console')
bullet('Archived Tab')
bullet('Settings')

heading3('5.2.1 Active Console')
para('This displays the below tabs: severity, tracker, call to action/human in loop, alerts timestamp, status, description. Except tracker and description, rest others are non-clickable. Also displays various filters as below:')
bold_bullet('Severity: ', 'Drop-down menu containing the level of severity — Low, High, Medium, Critical and also All in a single view.')
bold_bullet('Anomaly Type: ', 'Drop-down menu which displays Monitor and track onsite vehicle time log, watch list person detected, weapon detected.')
bold_bullet('Space: ', 'Dropdown — displays from which side of the camera the user wants to check the details.')
bold_bullet('Action Taken: ', 'Dropdown — displays if information sent to 911, Email, SMS or Rapid SOS.')
bold_bullet('Search: ', 'General search where the user can search based on his knowledge and experience.')

para('Tracker Field:', bold=True)
para('This is a clickable field which upon click displays:')
p = doc.add_paragraph()
r = p.add_run('Alert Details:')
r.bold = True
r.font.size = Pt(10.5)
bullet('Event timestamp, Description, Space details, Severity, Status', 1)
bullet('Change status dropdown (reviewed, false positive, completed notes, audited notes, action, action taken)', 1)
bullet('Camera footage — can be expanded', 1)
bullet('Action taken comments, Save / Cancel', 1)

p = doc.add_paragraph()
r = p.add_run('Track Incident Tab:')
r.bold = True
r.font.size = Pt(10.5)
bold_bullet('Action Taken Tab: ', 'Alert ID, severity, location, response timeline / Audit trail / log.', 1)
bold_bullet('Notification Sent Tab: ', 'List of users notified with date, timestamp, and status (delivered, queued).', 1)

para('Description Field:', bold=True)
para('This is a clickable field which upon click displays the same Alert Details and Track Incident information as the Tracker field above.')

heading3('5.2.2 Archived Tab')
para('Displays tabs: severity, timestamp, status and description (clickable). Same filters as Active Console (Severity, Anomaly Type, Space, Action Taken, Search). Description field shows same Alert Details and Track Incident information.')

heading3('5.2.3 Settings')
para('Settings available based on roles assigned to a user. Below are the settings available for an Admin:')

bold_bullet('Admin Settings: ', 'Admin can sort settings with the dropdown menu — By severity and By created date.')
bold_bullet('Escalation Settings: ', 'Admin can create rules for escalation based on the below factors:')
bullet('Rules: Type of escalations (Brandished guns & knives, guns and knives, vehicle time log, watchlist person detected).', 1)
bullet('When Brandished guns & knives and watchlist person detected are selected, the rule displays: Types of Levels, Escalated to, Pick actions and notifications. Levels can be up to 3. Auto expire in 5 min checkbox available.', 1)
bullet('Below the rules tab, displays the list of rules created.', 1)

bold_bullet('Alert Assignment: ', 'Admin can create a rule to whom the alerts can be sent.')
bullet('Brandished Guns & Knives — Super Admin and Security Officer', 1)
bullet('Guns & Knives 123 — NA', 1)
bullet('Vehicle Time Log — Security Admin', 1)
bullet('Watchlist Person Detected — Super Admin and Security Officer', 1)

heading2('5.3 Human in Loop Console')
para('A human has to verify if the alerts are true or false. This tab contains:')

bold_bullet('Alerts Console: ', 'List of alerts with severity, tracker, timestamp, status, description. Filters: severity, anomaly type, space, action taken, and search box.')
bold_bullet('Performance Analytics: ', 'User can review the performance analytics of the reviewers selected from the dropdown and the time frame. This displays: total alerts, how many reviewed, how many are false, how many are weapon detected, how many are escalated, average response time. This report can also be exported.')

add_horizontal_line()

# ═══════════════════════════════════════════════════════
# MODULE 6 – CONFIGURATIONS
# ═══════════════════════════════════════════════════════
heading1('6. Configurations')
para('Here, the admin can configure Schools/Spaces, cameras, and sensors.')

heading2('6.1 School/Space Configuration')
para('[Configuration details as per system requirements]')

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# GLOSSARY
# ═══════════════════════════════════════════════════════
heading1('Glossary')
glossary = [
    ('POI', 'Person of Interest'),
    ('OOI', 'Object of Interest'),
    ('OTP', 'One-Time Password'),
    ('BDD', 'Behavior-Driven Development'),
    ('RTSP', 'Real-Time Streaming Protocol'),
    ('FRD', 'Functional Requirements Document'),
]
t = doc.add_table(rows=len(glossary)+1, cols=2)
t.style = 'Light List Accent 1'
t.rows[0].cells[0].paragraphs[0].add_run('Acronym').bold = True
t.rows[0].cells[1].paragraphs[0].add_run('Definition').bold = True
for i, (acronym, defn) in enumerate(glossary, 1):
    t.rows[i].cells[0].paragraphs[0].add_run(acronym).font.size = Pt(10)
    t.rows[i].cells[1].paragraphs[0].add_run(defn).font.size = Pt(10)

# ── SAVE ──
doc.save(r'D:\Satish\Playwright\SAFE SPACE for School.docx')
print('Enhanced document saved successfully!')
